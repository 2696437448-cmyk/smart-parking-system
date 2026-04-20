from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path("/Users/yanchen/VscodeProject/smart-parking-thesis")
THESIS_DIR = BASE_DIR / "docs" / "thesis-docs"
MD_PATH = THESIS_DIR / "thesis-draft.md"
ASSET_DIR = THESIS_DIR / "assets" / "generated-final"
DOCX_PATH = THESIS_DIR / "智慧社区多源数据融合的车位动态调度与共享系统设计_最终版.docx"

CN_FONT_PATH = Path("/System/Library/Fonts/Supplemental/Songti.ttc")
EN_FONT_PATH = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = EN_FONT_PATH if bold and not CN_FONT_PATH.exists() else CN_FONT_PATH
    return ImageFont.truetype(str(path), size=size)


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, *, fill: str, outline: str) -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
    font = load_font(28)
    x1, y1, x2, y2 = box
    center_x = (x1 + x2) // 2
    center_y = (y1 + y2) // 2
    lines = text.split("\n")
    total_height = sum(font.getbbox(line)[3] - font.getbbox(line)[1] for line in lines) + (len(lines) - 1) * 6
    current_y = center_y - total_height // 2
    for line in lines:
        bbox = font.getbbox(line)
        line_w = bbox[2] - bbox[0]
        line_h = bbox[3] - bbox[1]
        draw.text((center_x - line_w // 2, current_y), line, fill="#111827", font=font)
        current_y += line_h + 6


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, color: str = "#111827", width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if x1 == x2 and y1 == y2:
        return
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        tip = (x2, y2)
        wing1 = (x2 - 18 * direction, y2 - 10)
        wing2 = (x2 - 18 * direction, y2 + 10)
    else:
        direction = 1 if y2 > y1 else -1
        tip = (x2, y2)
        wing1 = (x2 - 10, y2 - 18 * direction)
        wing2 = (x2 + 10, y2 - 18 * direction)
    draw.polygon([tip, wing1, wing2], fill=color)


def draw_title(draw: ImageDraw.ImageDraw, text: str, width: int) -> None:
    font = load_font(36)
    bbox = font.getbbox(text)
    text_w = bbox[2] - bbox[0]
    draw.text(((width - text_w) / 2, 26), text, fill="#111827", font=font)


def new_canvas(width: int = 1600, height: int = 900, title: str | None = None) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    if title:
        draw_title(draw, title, width)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    path = ASSET_DIR / name
    img.save(path)
    return path


def build_architecture() -> Path:
    img, draw = new_canvas(title="系统总体架构图")
    draw_box(draw, (110, 190, 360, 320), "业主端用户", fill="#dbeafe", outline="#2563eb")
    draw_box(draw, (1240, 190, 1490, 320), "物业端用户", fill="#dbeafe", outline="#2563eb")
    draw_box(draw, (545, 150, 1055, 300), "Frontend App\nVue3 + TypeScript + Pinia", fill="#dcfce7", outline="#16a34a")
    draw_box(draw, (590, 360, 1010, 500), "Gateway Service\nSpring Cloud Gateway", fill="#fef3c7", outline="#d97706")
    draw_box(draw, (120, 610, 460, 760), "Parking Service\n预约 / 计费 / 订单 / Dashboard", fill="#fee2e2", outline="#dc2626")
    draw_box(draw, (630, 610, 970, 760), "Model Service\nLSTM-Lite / Hungarian", fill="#ede9fe", outline="#7c3aed")
    draw_box(draw, (1140, 610, 1480, 760), "Realtime Service\nWebSocket / Polling Fallback", fill="#e0f2fe", outline="#0284c7")
    draw_box(draw, (340, 800, 520, 870), "MySQL", fill="#f3f4f6", outline="#4b5563")
    draw_box(draw, (710, 800, 890, 870), "Redis", fill="#f3f4f6", outline="#4b5563")
    draw_box(draw, (1080, 800, 1260, 870), "RabbitMQ", fill="#f3f4f6", outline="#4b5563")

    draw_arrow(draw, (360, 255), (545, 225))
    draw_arrow(draw, (1240, 255), (1055, 225))
    draw_arrow(draw, (800, 300), (800, 360))
    draw_arrow(draw, (800, 500), (290, 610))
    draw_arrow(draw, (800, 500), (800, 610))
    draw_arrow(draw, (800, 500), (1310, 610))
    draw_arrow(draw, (290, 760), (430, 800))
    draw_arrow(draw, (800, 760), (800, 800))
    draw_arrow(draw, (1310, 760), (1170, 800))
    return save(img, "fig_3_1_architecture.png")


def build_dataflow() -> Path:
    img, draw = new_canvas(title="主数据流与业务流示意图")
    nodes = [
        ((420, 120, 1180, 210), "sensor_event_raw / lpr_event_raw /\nresident_trip_raw", "#dbeafe", "#2563eb"),
        ((500, 250, 1100, 340), "ETL与特征处理", "#dcfce7", "#16a34a"),
        ((420, 380, 1180, 470), "forecast_feature_table /\ndispatch_input_table", "#fef3c7", "#d97706"),
        ((540, 510, 1060, 600), "Model Service\n预测与调度", "#ede9fe", "#7c3aed"),
        ((540, 640, 1060, 730), "Parking Service\n预约 / 计费 / 订单", "#fee2e2", "#dc2626"),
    ]
    for box, text, fill, outline in nodes:
        draw_box(draw, box, text, fill=fill, outline=outline)
    draw_box(draw, (80, 620, 360, 710), "Owner Dashboard", fill="#e0f2fe", outline="#0284c7")
    draw_box(draw, (1240, 620, 1520, 710), "Navigation", fill="#e0f2fe", outline="#0284c7")
    draw_box(draw, (200, 780, 600, 870), "Billing Records / Revenue Summary", fill="#f3f4f6", outline="#4b5563")
    draw_box(draw, (1000, 780, 1400, 870), "Admin Dashboard", fill="#e0f2fe", outline="#0284c7")
    for start, end in [((800, 210), (800, 250)), ((800, 340), (800, 380)), ((800, 470), (800, 510)), ((800, 600), (800, 640))]:
        draw_arrow(draw, start, end)
    draw_arrow(draw, (540, 685), (360, 665))
    draw_arrow(draw, (1060, 685), (1240, 665))
    draw_arrow(draw, (620, 730), (460, 780))
    draw_arrow(draw, (1060, 555), (1200, 780))
    draw_arrow(draw, (600, 825), (1000, 825))
    return save(img, "fig_3_2_dataflow.png")


def build_predict_dispatch() -> Path:
    img, draw = new_canvas(height=620, title="预测与调度协同流程图")
    boxes = [
        ((70, 230, 320, 360), "多源特征数据", "#dbeafe", "#2563eb"),
        ((360, 230, 610, 360), "LSTM-Lite预测", "#dcfce7", "#16a34a"),
        ((650, 230, 900, 360), "供需缺口 /\n占用趋势", "#fef3c7", "#d97706"),
        ((940, 230, 1190, 360), "Hungarian调度", "#ede9fe", "#7c3aed"),
        ((1230, 230, 1530, 360), "候选车位匹配结果", "#fee2e2", "#dc2626"),
        ((580, 430, 1020, 550), "Parking Service / 推荐结果 /\n预约主链 / 经营图表", "#e0f2fe", "#0284c7"),
    ]
    for box, text, fill, outline in boxes:
        draw_box(draw, box, text, fill=fill, outline=outline)
    draw_arrow(draw, (320, 295), (360, 295))
    draw_arrow(draw, (610, 295), (650, 295))
    draw_arrow(draw, (900, 295), (940, 295))
    draw_arrow(draw, (1190, 295), (1230, 295))
    draw_arrow(draw, (775, 360), (775, 430))
    draw_arrow(draw, (1380, 360), (920, 430))
    return save(img, "fig_4_1_predict_dispatch.png")


def build_sequence() -> Path:
    img, draw = new_canvas(height=980, title="业主预约主链时序图")
    participants = ["业主端页面", "Gateway", "Parking Service", "Redis/Lock", "MySQL", "Model Service"]
    xs = [160, 400, 640, 880, 1120, 1360]
    top_y = 130
    for x, name in zip(xs, participants):
        draw_box(draw, (x - 85, top_y, x + 85, top_y + 60), name, fill="#f3f4f6", outline="#4b5563")
        draw.line([(x, top_y + 60), (x, 870)], fill="#9ca3af", width=2)
    messages = [
        (0, 1, 220, "提交预约请求"),
        (1, 2, 290, "转发请求并透传trace"),
        (2, 3, 360, "幂等校验与锁控制"),
        (2, 4, 430, "校验车位与订单状态"),
        (2, 5, 500, "获取调度或推荐参考"),
        (5, 2, 570, "返回匹配结果"),
        (2, 4, 640, "写入订单与计费信息"),
        (2, 1, 710, "返回预约结果"),
        (1, 0, 780, "展示订单、账单与导航入口"),
    ]
    font = load_font(22)
    for src, dst, y, label in messages:
        start = (xs[src], y)
        end = (xs[dst], y)
        draw_arrow(draw, start, end, width=3)
        bbox = font.getbbox(label)
        text_w = bbox[2] - bbox[0]
        mid_x = (start[0] + end[0]) / 2
        draw.text((mid_x - text_w / 2, y - 26), label, fill="#111827", font=font)
    return save(img, "fig_4_2_sequence.png")


def build_admin_flow() -> Path:
    img, draw = new_canvas(height=860, title="物业端Dashboard数据编排与状态回退流程图")
    nodes = [
        ((580, 110, 1020, 190), "Admin页面进入", "#dbeafe", "#2563eb"),
        ((580, 230, 1020, 310), "useAdminDashboardView", "#dcfce7", "#16a34a"),
        ((580, 350, 1020, 430), "GET /api/v1/admin/dashboard", "#fef3c7", "#d97706"),
        ((220, 490, 680, 570), "实时状态联动", "#ede9fe", "#7c3aed"),
        ((920, 490, 1380, 570), "图表数据转换", "#fee2e2", "#dc2626"),
        ((220, 660, 680, 740), "degraded / stale 提示", "#fef2f2", "#b91c1c"),
        ((920, 660, 1380, 740), "正常展示", "#e0f2fe", "#0284c7"),
        ((220, 760, 680, 840), "Polling回退", "#f3f4f6", "#4b5563"),
    ]
    for box, text, fill, outline in nodes:
        draw_box(draw, box, text, fill=fill, outline=outline)
    for start, end in [((800, 190), (800, 230)), ((800, 310), (800, 350)), ((800, 430), (450, 490)), ((800, 430), (1150, 490))]:
        draw_arrow(draw, start, end)
    draw_arrow(draw, (450, 570), (450, 660))
    draw_arrow(draw, (1150, 570), (1150, 660))
    draw_arrow(draw, (450, 740), (450, 760))
    draw_arrow(draw, (680, 800), (920, 700))
    return save(img, "fig_4_3_admin_flow.png")


def build_state() -> Path:
    img, draw = new_canvas(width=1300, height=620, title="页面状态统一表达示意图")
    boxes = [
        ((90, 260, 250, 340), "loading", "#dbeafe", "#2563eb"),
        ((320, 120, 480, 200), "ready", "#dcfce7", "#16a34a"),
        ((320, 260, 480, 340), "empty", "#fef3c7", "#d97706"),
        ((320, 400, 480, 480), "error", "#fee2e2", "#dc2626"),
        ((560, 120, 760, 200), "degraded", "#ede9fe", "#7c3aed"),
        ((860, 120, 1040, 200), "stale", "#e0f2fe", "#0284c7"),
    ]
    for box, text, fill, outline in boxes:
        draw_box(draw, box, text, fill=fill, outline=outline)
    draw_arrow(draw, (250, 300), (320, 160))
    draw_arrow(draw, (250, 300), (320, 300))
    draw_arrow(draw, (250, 300), (320, 440))
    draw_arrow(draw, (480, 160), (560, 160))
    draw_arrow(draw, (760, 160), (860, 160))
    draw_arrow(draw, (950, 200), (430, 260))
    draw_arrow(draw, (400, 400), (180, 340))
    return save(img, "fig_5_1_state.png")


def build_pie() -> Path:
    img, draw = new_canvas(width=1200, height=760, title="Step40综合验收覆盖维度示意图")
    center = (350, 410)
    radius = 220
    values = [3, 1, 1, 1]
    labels = ["历史基线兼容", "Dashboard合同与状态表达", "聚合层模块化", "性能与默认完成态升级"]
    colors = ["#2563eb", "#16a34a", "#d97706", "#7c3aed"]
    total = sum(values)
    start = -90.0
    legend_font = load_font(24)
    for value, color in zip(values, colors):
        end = start + value / total * 360
        draw.pieslice((center[0] - radius, center[1] - radius, center[0] + radius, center[1] + radius), start, end, fill=color, outline="white")
        start = end
    for idx, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = 220 + idx * 90
        draw.rounded_rectangle((700, y, 740, y + 40), radius=6, fill=color, outline=color)
        text = f"{label}（{value}项）"
        draw.text((760, y + 4), text, fill="#111827", font=legend_font)
    return save(img, "fig_5_2_pie.png")


def generate_figures() -> dict[str, Path]:
    ensure_dirs()
    return {
        "图3-1": build_architecture(),
        "图3-2": build_dataflow(),
        "图4-1": build_predict_dispatch(),
        "图4-2": build_sequence(),
        "图4-3": build_admin_flow(),
        "图5-1": build_state(),
        "图5-2": build_pie(),
    }


def set_run_font(run, size: int = 12, *, bold: bool = False, center: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if center:
        run.italic = False


def set_paragraph_format(paragraph, *, first_line_indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0)


def add_text_paragraph(doc: Document, text: str, *, size: int = 12, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=indent, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_center_title(doc: Document, text: str, *, size: int = 18, bold: bool = True, space_after: int = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=15 if level == 1 else 13, bold=True)


def add_caption(doc: Document, text: str, *, above: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6 if above else 0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=True)


def add_toc(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = "右键目录并选择“更新域”后显示页码。"
    run.append(txt)
    fld.append(run)
    paragraph._p.append(fld)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(3.8)
    section.bottom_margin = Cm(3.8)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def extract_meta(lines: list[str]) -> dict[str, str]:
    wanted = ["中文题目", "英文题目", "学院（部）", "专业班级", "学生姓名", "指导教师", "完成日期"]
    meta: dict[str, str] = {}
    for key in wanted:
        for idx, line in enumerate(lines):
            if line.strip() == f"## {key}":
                for value in lines[idx + 1 : idx + 5]:
                    if value.strip() and not value.startswith("#"):
                        meta[key] = value.strip().strip("`")
                        break
                break
    return meta


def add_cover(doc: Document, meta: dict[str, str]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_center_title(doc, "[学校名称]本科毕业设计说明书", size=20)
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, meta["中文题目"], size=18)
    add_center_title(doc, meta["英文题目"], size=14, bold=False)
    for _ in range(3):
        doc.add_paragraph()
    fields = [
        ("学院（部）", meta["学院（部）"]),
        ("专业班级", meta["专业班级"]),
        ("学生姓名", meta["学生姓名"]),
        ("指导教师", meta["指导教师"]),
        ("完成日期", meta["完成日期"]),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, size=14)
    doc.add_page_break()


def is_table_line(text: str) -> bool:
    return text.strip().startswith("|") and text.strip().endswith("|")


def split_table_row(text: str) -> list[str]:
    return [cell.strip() for cell in text.strip().strip("|").split("|")]


def build_word_table(doc: Document, caption: str | None, rows: list[list[str]]) -> None:
    if caption:
        add_caption(doc, caption, above=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=10, bold=(r_idx == 0))
    doc.add_paragraph()


def gather_paragraph(lines: list[str], start: int) -> tuple[str, int]:
    collected = [lines[start].strip()]
    idx = start + 1
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            break
        if line.startswith("#") or line.startswith("图") or line.startswith("表") or line.startswith("```") or is_table_line(line):
            break
        collected.append(line)
        idx += 1
    return "\n".join(collected), idx


def add_markdown_content(doc: Document, lines: list[str], figures: dict[str, Path], meta: dict[str, str]) -> None:
    start = lines.index("## 摘要")
    idx = start
    pending_table_caption: str | None = None
    pending_figure_caption: str | None = None
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue

        if stripped == "## 摘要":
            add_center_title(doc, meta["中文题目"], size=16)
            add_center_title(doc, "摘要", size=15)
            idx += 1
            continue
        if stripped == "## ABSTRACT":
            doc.add_page_break()
            add_center_title(doc, meta["英文题目"], size=15, bold=False)
            add_center_title(doc, "ABSTRACT", size=15)
            idx += 1
            continue
        if stripped == "## 目录":
            doc.add_page_break()
            add_center_title(doc, "目录", size=15)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_toc(p)
            doc.add_page_break()
            idx += 1
            while idx < len(lines) and not re.match(r"## \d", lines[idx].strip()):
                idx += 1
            continue
        if stripped in {"## 参考文献", "## 致谢"}:
            doc.add_page_break()
            add_heading(doc, stripped[3:], 1)
            idx += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            idx += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
            idx += 1
            continue
        if re.match(r"^表\d+-\d+\s", stripped):
            pending_table_caption = stripped
            idx += 1
            continue
        if re.match(r"^图\d+-\d+\s", stripped):
            pending_figure_caption = stripped
            idx += 1
            continue
        if stripped.startswith("```"):
            fence_name = stripped.strip("`")
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                idx += 1
            if pending_figure_caption and pending_figure_caption[:4] in figures:
                figure_key = pending_figure_caption[:4]
                doc.add_picture(str(figures[figure_key]), width=Inches(6.2))
                add_caption(doc, pending_figure_caption, above=False)
                pending_figure_caption = None
            idx += 1
            continue
        if is_table_line(stripped):
            table_lines: list[list[str]] = []
            while idx < len(lines) and is_table_line(lines[idx].strip()):
                row = split_table_row(lines[idx])
                table_lines.append(row)
                idx += 1
            if len(table_lines) >= 2 and all(cell.startswith("---") or cell.endswith("---") or set(cell) <= {"-", ":"} for cell in table_lines[1]):
                table_lines.pop(1)
            build_word_table(doc, pending_table_caption, table_lines)
            pending_table_caption = None
            continue
        if stripped.startswith("资料来源："):
            add_text_paragraph(doc, stripped, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
            idx += 1
            continue
        paragraph_text, new_idx = gather_paragraph(lines, idx)
        add_text_paragraph(doc, paragraph_text.replace("\n", " "), size=12)
        idx = new_idx


def build_docx() -> Path:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    meta = extract_meta(lines)
    figures = generate_figures()
    doc = Document()
    configure_document(doc)
    add_cover(doc, meta)
    add_markdown_content(doc, lines, figures, meta)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    output = build_docx()
    print(output)
